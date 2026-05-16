"""
Script to read Word documents and extract requirements for implementation
"""
from docx import Document
import os

def read_word_file(file_path):
    """Read and extract text from a Word document"""
    try:
        doc = Document(file_path)
        content = []

        print(f"\n{'='*80}")
        print(f"Reading: {os.path.basename(file_path)}")
        print(f"{'='*80}\n")

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                content.append(para.text)
                print(para.text)

        # Also extract tables if any
        if doc.tables:
            print("\n--- TABLES FOUND ---\n")
            for table_idx, table in enumerate(doc.tables):
                print(f"\nTable {table_idx + 1}:")
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    print(" | ".join(row_text))

        return content

    except (RuntimeError, ValueError, TypeError, AttributeError, OSError) as e:
        print(f"Error reading {file_path}: {str(e)}")
        return []

if __name__ == "__main__":
    # Find all .docx files in current directory
    word_files = [f for f in os.listdir('.') if f.endswith('.docx')]

    all_content = {}
    for word_file in word_files:
        content = read_word_file(word_file)
        all_content[word_file] = content

    print(f"\n\n{'='*80}")
    print(f"Found {len(word_files)} Word document(s)")
    print(f"{'='*80}")

